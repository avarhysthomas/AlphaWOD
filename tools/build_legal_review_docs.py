from __future__ import annotations

import argparse
import hashlib
import json
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


OUTPUT_DIR = Path("docs/legal-review/2026-08-23")
VERSION_DATE = "23 August 2026"
ENTITY = "ZERO ALPHA FITNESS LTD"
COMPANY_NUMBER = "15978998"
TRADING_NAME = "Zero Alpha Fitness"
ADDRESS = "Unit 3, Felinfoel Business Hub, Llanelli, SA14 8BE"
REGISTERED_OFFICE = "18 Bryngwyn Bach, Llanelli, United Kingdom, SA14 8SH"
SUPPORT_EMAIL = "support@zeroalphafitness.co.uk"
SITE_URL = "https://alpha-wod.vercel.app/"

NAVY = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1E293B"
MUTED = "5F6B7A"
LIGHT_BLUE = "EAF2F8"
PALE_BLUE = "F4F8FC"
PALE_GREY = "F2F4F7"
BORDER = "C9D4DF"
RED = "A61B1B"
PALE_RED = "FDECEC"
AMBER = "8A4B08"
PALE_AMBER = "FFF4DF"
GREEN = "22633A"
WHITE = "FFFFFF"


@dataclass(frozen=True)
class DocSpec:
    registry_key: str
    filename: str
    title: str
    subtitle: str
    doc_id: str
    subject: str


SPECS = {
    "terms": DocSpec(
        "membershipTerms",
        "01-membership-terms.docx",
        "Membership Terms",
        "Public membership purchase and ongoing membership",
        "ZAF-TERMS-2026-08-23-01",
        "Membership terms for public membership purchase",
    ),
    "privacy": DocSpec(
        "privacyNotice",
        "02-privacy-notice.docx",
        "Privacy Notice",
        "Membership, payment, AlphaWOD and participant information",
        "ZAF-PRIVACY-2026-08-23-01",
        "Privacy notice for membership and AlphaWOD data",
    ),
    "cancel": DocSpec(
        "cancellationPolicy",
        "03-cancellation-refund-cooling-off-policy.docx",
        "Cancellation, Refund and Cooling-off Policy",
        "Plain-English rules for rolling monthly memberships",
        "ZAF-CANCEL-2026-08-23-01",
        "Cancellation, refund and cooling-off policy",
    ),
    "adult": DocSpec(
        "adultWaiver",
        "04-adult-participant-waiver.docx",
        "Adult Participant Waiver and Risk Acknowledgement",
        "For every participant aged 18 or over",
        "ZAF-ADULT-WAIVER-2026-08-23-01",
        "Adult participant waiver and risk acknowledgement",
    ),
    "guardian": DocSpec(
        "guardianAddendum",
        "05-parent-guardian-addendum.docx",
        "Parent/Guardian Consent and Youth Membership Addendum",
        "For Youngstars (ages 6–11) and Teenstars (ages 12–16), including multi-child checkouts",
        "ZAF-GUARDIAN-2026-08-23-01",
        "Parent and guardian consent and youth membership addendum",
    ),
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in edges:
            continue
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edges[edge].items():
            element.set(qn("w:" + key), str(value))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_hyperlink(paragraph, text: str, url: str, *, font_size: float | None = None) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.extend([color, underline])
    if font_size is not None:
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), str(int(font_size * 2)))
        size_cs = OxmlElement("w:szCs")
        size_cs.set(qn("w:val"), str(int(font_size * 2)))
        run_pr.extend([size, size_cs])
    run.append(run_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_link_style(run) -> None:
    run.font.color.rgb = RGBColor.from_string(BLUE)
    run.font.underline = True


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def configure_document(doc: Document, spec: DocSpec) -> None:
    # Use one header/footer definition throughout. Separate odd/even definitions
    # render inconsistently in some LibreOffice versions and can push the even-
    # page footer fields outside the printable area.
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.widow_control = True

    for style_name, size, color, before, after in (
        ("Title", 24, NAVY, 0, 4),
        ("Subtitle", 11.5, MUTED, 0, 12),
        ("Heading 1", 16, BLUE, 16, 7),
        ("Heading 2", 13, DARK_BLUE, 12, 5),
        ("Heading 3", 11.5, DARK_BLUE, 9, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.left_indent = Inches(0.28)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.08
        style.paragraph_format.widow_control = True

    if "Small Print" not in styles:
        small = styles.add_style("Small Print", WD_STYLE_TYPE.PARAGRAPH)
    else:
        small = styles["Small Print"]
    small.font.name = "Calibri"
    small.font.size = Pt(8.5)
    small.font.color.rgb = RGBColor.from_string(MUTED)
    small.paragraph_format.space_after = Pt(3)
    small.paragraph_format.line_spacing = 1.0

    props = doc.core_properties
    props.title = spec.title
    props.subject = spec.subject
    props.author = ENTITY
    props.keywords = "membership, customer terms, Zero Alpha Fitness"
    props.comments = "Final customer-facing publication copy."
    props.created = datetime(2026, 8, 23, 12, 0, tzinfo=timezone.utc)
    props.modified = datetime(2026, 8, 23, 12, 0, tzinfo=timezone.utc)

    def format_header(header) -> None:
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.clear()
        hp.paragraph_format.space_after = Pt(0)

    def format_footer(footer) -> None:
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)
        left = fp.add_run(spec.doc_id)
        left.font.name = "Calibri"
        left.font.size = Pt(8)
        left.font.color.rgb = RGBColor.from_string(MUTED)
        tab_stops = fp.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(5.55))
        mid = fp.add_run("\tPage ")
        mid.font.size = Pt(8)
        mid.font.color.rgb = RGBColor.from_string(MUTED)
        add_field(fp, "PAGE")
        of_run = fp.add_run(" of ")
        of_run.font.size = Pt(8)
        of_run.font.color.rgb = RGBColor.from_string(MUTED)
        add_field(fp, "NUMPAGES")

    format_header(section.header)
    format_footer(section.footer)


def add_document_title(doc: Document, spec: DocSpec) -> None:
    p = doc.add_paragraph(style="Title")
    p.add_run(spec.title)
    p = doc.add_paragraph(style="Subtitle")
    p.add_run(spec.subtitle)

    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [Inches(2.05), Inches(2.05), Inches(2.05)]
    labels = [
        ("VERSION DATE", VERSION_DATE),
        ("DOCUMENT ID", spec.doc_id),
        ("CONTRACTING ENTITY", f"{ENTITY} · {COMPANY_NUMBER}"),
        ("TRADING NAME", TRADING_NAME),
        ("PUBLIC CONTACT", SUPPORT_EMAIL),
        ("WEBSITE", SITE_URL),
    ]
    for idx, cell in enumerate(table._cells):
        cell.width = widths[idx % 3]
        set_cell_shading(cell, PALE_GREY)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "4", "color": WHITE},
            bottom={"val": "single", "sz": "4", "color": WHITE},
            left={"val": "single", "sz": "4", "color": WHITE},
            right={"val": "single", "sz": "4", "color": WHITE},
        )
        set_cell_margins(cell, 75, 100, 75, 100)
        label, value = labels[idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(label + "\n")
        r.font.size = Pt(7.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(MUTED)
        r2 = p.add_run(value)
        r2.font.size = Pt(8.6)
        r2.font.color.rgb = RGBColor.from_string(INK)
    for row in table.rows:
        prevent_row_split(row)


def add_draft_banner(doc: Document) -> None:
    """Retained as a no-op so the final builder cannot emit review banners."""
    del doc


def add_callout(doc: Document, title: str, body: str, kind: str = "info") -> None:
    palette = {
        "info": (BLUE, LIGHT_BLUE),
        "warning": (AMBER, PALE_AMBER),
        "danger": (RED, PALE_RED),
        "success": (GREEN, "EAF6EE"),
    }
    color, fill = palette[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.55)
    set_cell_shading(cell, fill)
    set_cell_border(cell, left={"val": "single", "sz": "18", "color": color})
    set_cell_margins(cell, 110, 150, 105, 150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(color)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.04
    for r2 in p2.runs:
        r2.font.size = Pt(9.5)
    prevent_row_split(table.rows[0])
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)
    after.paragraph_format.space_before = Pt(0)


def add_para(doc: Document, text: str, *, bold_lead: str | None = None, style: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        r1.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_quote(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    cell.width = Inches(6.45)
    set_cell_shading(cell, PALE_BLUE)
    set_cell_border(cell, left={"val": "single", "sz": "12", "color": BLUE})
    set_cell_margins(cell, 95, 140, 95, 140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    r.font.italic = True
    r.font.size = Pt(10)
    prevent_row_split(table.rows[0])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def add_source_list(
    doc: Document,
    sources: Sequence[tuple[str, str]],
    *,
    compact_columns: bool = False,
) -> None:
    if not compact_columns:
        for label, url in sources:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.keep_together = True
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            add_hyperlink(p, label, url, font_size=8.5)
        return

    table = doc.add_table(rows=(len(sources) + 1) // 2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, (label, url) in enumerate(sources):
        cell = table.cell(index // 2, index % 2)
        set_cell_margins(cell, 15, 80, 15, 80)
        p = cell.paragraphs[0]
        p.style = doc.styles["List Bullet"]
        p.paragraph_format.keep_together = True
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        add_hyperlink(p, label, url, font_size=7.5)
    for row in table.rows:
        prevent_row_split(row)


def add_review_appendix(
    doc: Document,
    review_points: Sequence[str],
    sources: Sequence[tuple[str, str]],
    *,
    high_priority: str | None = None,
    compact: bool = False,
) -> None:
    """Retained as a no-op so review-only material cannot enter final copies."""
    del doc, review_points, sources, high_priority, compact


def add_acceptance_block(doc: Document, heading: str, statements: Sequence[str], signer: str) -> None:
    doc.add_heading(heading, level=2)
    add_para(
        doc,
        "The checkout should show each required statement beside its own unticked control. Acceptance must be affirmative; do not pre-tick, infer or bundle optional marketing consent.",
    )
    for statement in statements:
        add_quote(doc, "☐ " + statement)
    add_para(
        doc,
        f"Evidence to retain: {signer}; the exact document version(s); the exact statement(s) displayed; typed name where required; UTC timestamp plus Europe/London display time; authenticated account or verified-email context; Stripe Checkout/subscription identifiers where relevant; and only the technical audit data disclosed in the Privacy Notice.",
    )


def build_terms() -> Path:
    spec = SPECS["terms"]
    doc = Document()
    configure_document(doc, spec)
    add_document_title(doc, spec)
    add_draft_banner(doc)

    doc.add_heading("1. About these Terms", level=1)
    add_para(
        doc,
        f"These Membership Terms form the contract for a Zero Alpha Fitness membership purchased online from {ENTITY}, a company registered in England and Wales under company number {COMPANY_NUMBER}, trading as {TRADING_NAME} (we, us or our). Our registered office is {REGISTERED_OFFICE}. Our public trading and contact address is {ADDRESS}. Contact us at {SUPPORT_EMAIL}.",
    )
    add_para(
        doc,
        "For an adult membership, the named participant must buy the membership for themselves and is also the payer and signer. For one or more participants under 18, the payer and signer must be their parent or legal guardian, or another adult with lawful authority to enrol every named child.",
    )
    add_para(
        doc,
        "These Terms should be read with the Cancellation, Refund and Cooling-off Policy, the Privacy Notice, and the applicable Adult Participant Waiver or Parent/Guardian Consent and Youth Membership Addendum. Those documents are presented before purchase and form part of the membership arrangement where stated.",
    )

    doc.add_heading("2. Eligibility and who must accept what", level=1)
    add_bullets(
        doc,
        [
            "An adult participant must be aged 18 or over and must personally accept the Adult Participant Waiver and Risk Acknowledgement before taking part.",
            "A Youngstars participant must be aged 6 to 11 inclusive. A Teenstars participant must be aged 12 to 16 inclusive. Every child named in a checkout must meet the selected plan’s age band.",
            "A youth participant’s guardian must be the payer, confirm their relationship and lawful authority for each named child, and accept the Parent/Guardian Consent and Youth Membership Addendum for each child.",
            "An adult participant accepts these Terms, the payment obligation and the Adult Participant Waiver for their own membership. Third-party purchase of an adult membership is not supported.",
            "You must give accurate, complete and current information. You must not buy a youth membership for a child outside the stated age range or misrepresent authority to act for another person.",
        ],
    )

    doc.add_heading("3. Membership options and prices", level=1)
    add_para(doc, "The initial public catalogue is:")
    add_bullets(
        doc,
        [
            "Adult Unlimited Membership — £60 per month. This is the only paid membership that grants eligible AlphaWOD access after the first required payment succeeds.",
            "Adult Ladies Only Membership — £50 per month.",
            "Adult Gym Only — £45 per month.",
            "Youth Membership — Youngstars (ages 6–11) at £30 per child per month, or Teenstars (ages 12–16) at £35 per child per month.",
        ],
    )
    add_para(
        doc,
        "There is no joining fee, free trial or minimum term. Each membership is a rolling monthly contract until cancelled under section 9. We are not currently VAT registered. The price displayed at checkout is the total customer price; no VAT invoice or VAT breakdown is offered. If our tax status changes, we will update the presentation and notices before applying any change.",
    )
    add_para(
        doc,
        "One youth Checkout may cover between one and ten children in the same selected plan. Youngstars and Teenstars cannot be combined in one Checkout. The contracted quantity is the number of children confirmed at Checkout or in a later written change confirmed by us.",
    )
    add_para(
        doc,
        "When the contracted quantity is two or more children, a recurring 15% multi-child discount applies to the full undiscounted monthly subtotal (the price per child multiplied by the contracted quantity). It does not apply at quantity one and does not combine with another promotion unless Stripe explicitly shows both before confirmation. Non-attendance does not change the contracted quantity, recurring total or discount.",
    )
    add_para(
        doc,
        "We may offer promotion codes subject to their stated eligibility, duration and limits. The automatic youth multi-child discount does not combine with another promotion unless Stripe explicitly shows both before confirmation. A promotion does not change the underlying rolling nature of the contract unless its terms expressly say so. A code has no cash value and cannot be applied retrospectively unless required by law or expressly agreed.",
    )

    doc.add_heading("4. How the contract is made", level=1)
    add_numbered(
        doc,
        [
            "Choose the membership and, for youth membership, one plan and the number of children in that plan.",
            "Provide the required details for every named participant and the payer or guardian, review the documents, and complete each required acceptance or signature.",
            "Review the amount due today, the first payment or billing date and, for a youth membership, the quantity, price per child, undiscounted subtotal, multi-child discount and resulting recurring total shown by Stripe. During the founding presale, Checkout must show £0 due today and the first monthly payment on 1 September 2026; after opening, Stripe may show an immediate prorated charge to the next first of the month.",
            "Complete Stripe Checkout. Available payment methods are those Stripe displays for that transaction and may vary by device, location, currency and eligibility.",
            "The contract is formed when Stripe confirms completion of Checkout and we issue an on-screen or email confirmation, unless we promptly tell you that a clear pricing, eligibility or technical error prevented acceptance and refund any amount taken.",
        ],
    )
    add_para(
        doc,
        "Stripe will present an unambiguous final confirmation control and summary. Before service begins, we will email a durable confirmation containing the agreed plan, every named participant, contracted quantity, price per child, undiscounted subtotal, discount, resulting recurring total, next payment date, accepted document versions, cancellation information and signed acceptance evidence; a changeable website link alone is not the durable copy. During the founding presale, nothing is charged today; membership starts and the first monthly payment is taken on 1 September 2026. For a youth membership, Checkout completion does not itself book a first session. We will contact the guardian by email to arrange onboarding and the first session.",
    )
    add_quote(
        doc,
        "Founding presale youth confirmation: You’re signed up. Nothing has been charged today. This membership starts and the first monthly payment is taken on 1 September 2026. Zero Alpha Fitness will contact you by email to arrange onboarding and the first session.",
    )

    doc.add_heading("5. Presale, billing date, proration and recurring authority", level=1)
    add_para(
        doc,
        "All memberships use the first day of each calendar month as the regular billing date, interpreted in Europe/London time. During the founding presale, nothing is charged today. Stripe securely saves the payment method; membership starts and the first monthly payment is taken on 1 September 2026. After opening, memberships start immediately and, if Checkout occurs after the first, Stripe calculates and displays an immediate prorated charge until the next first. The full monthly price is then charged on each first while membership continues.",
    )
    add_para(
        doc,
        "The amount Stripe displays before confirmation is authoritative for that Checkout. A presale Checkout must show £0 due today and a first payment date of 1 September 2026. We do not calculate a separate charge in the browser. If the displayed amount or billing date appears wrong, do not confirm; contact us first.",
    )
    add_para(
        doc,
        "By completing Checkout, the payer authorises Stripe and us to store the selected payment method as permitted by that payment method and to collect the amount shown today, which is £0 during the founding presale, and future recurring amounts without the payer being present. For a youth family subscription, that authority covers the displayed contracted quantity, price per child, subtotal, recurring 15% discount at quantity two or more, and resulting recurring total. The payer must keep the payment method valid and may update it through the secure Customer Portal.",
    )

    doc.add_heading("6. Price changes", level=1)
    add_para(
        doc,
        "We may change a recurring price only prospectively. We will give clear advance notice of the new amount and the date it would first apply, using the payer’s recorded email and any legally required method. The payer may cancel before the change takes effect. A price change will not alter amounts already paid or due for a completed billing event. Any legally required notice or cancellation right overrides this paragraph.",
    )

    doc.add_heading("7. What the membership provides", level=1)
    add_para(
        doc,
        "The selected membership provides access to the facilities, sessions or services described for that plan at the time of purchase, subject to opening hours, capacity, timetables, coaching instructions, reasonable safety rules and temporary closures. Each membership or youth place is personal to its named participant and cannot be sold, shared or transferred without our written agreement.",
    )
    add_para(
        doc,
        "We may make reasonable operational changes to timetables, instructors, equipment, class formats or facilities. We will not use this clause to remove the essential benefit of the membership without a fair remedy. If a change is material and adverse, we will give reasonable notice where practicable and any cancellation or refund rights required by law.",
    )
    add_para(
        doc,
        "Reaching the top of a youth age band does not automatically move a child, remove a place or change the recurring total. We will contact the payer about any transition. No plan, contracted-quantity, price or discount change takes effect until we confirm it in writing, including the effective date and resulting recurring total, and a new Checkout may be required.",
    )

    doc.add_heading("8. AlphaWOD access", level=1)
    add_bullets(
        doc,
        [
            "Adult Unlimited Membership qualifies the participant for AlphaWOD access only after the first required payment succeeds and while the subscription and payment status remain eligible.",
            "An existing AlphaWOD account holder should sign in and claim the purchase. Creating or paying for a duplicate active subscription may be blocked.",
            "Existing users whose AlphaWOD access was already approved before this purchase flow launches remain grandfathered unless their independent eligibility is later removed under a lawful policy.",
            "Administrators and SGPT staff may retain access through their role independently of a consumer membership.",
            "Youth memberships, Adult Ladies Only and Adult Gym Only do not automatically include AlphaWOD access.",
        ],
    )
    add_para(
        doc,
        "AlphaWOD access may be suspended or removed when the related entitlement is past due beyond the grace period, unpaid, fully refunded, subject to a lost payment dispute, cancelled or otherwise ineligible. A scheduled cancellation continues only through the paid access period. App availability also depends on compatible technology, internet service, security and maintenance.",
    )

    doc.add_heading("9. Ordinary cancellation of the rolling membership", level=1)
    add_callout(
        doc,
        "Prominent renewal rule",
        "To avoid the next first-of-month payment, a cancellation request must reach us at least 14 calendar days before that billing date. If it reaches us less than 14 days before the next first, that next monthly payment remains due and the membership ends at the end of the additional paid month.",
        kind="warning",
    )
    add_para(
        doc,
        "A payer may request cancellation through the signed-in cancellation-request flow or by emailing support@zeroalphafitness.co.uk. The online flow records and acknowledges the outcome automatically. Email requests are handled manually using the time the message reaches the support inbox. The request should identify the payer, every participant and the membership. We will send an acknowledgement showing the effective end date and any remaining scheduled charge.",
    )
    add_bullets(
        doc,
        [
            "If the request is received at least 14 calendar days before the next first, no payment is taken on that first and access ends at the end of the preceding day.",
            "If the request is received less than 14 calendar days before the next first, the payment on that first remains due; access continues for that paid month and ends at the end of the day before the following first.",
            "A cancellation request stops later renewals; it does not create a refund for time already supplied or an amount already due, except where the law requires one.",
            "Membership cannot be paused. A request to pause will not be treated as a cancellation unless the payer clearly asks to cancel.",
        ],
    )
    add_para(
        doc,
        "The statutory cooling-off right for a newly formed distance contract is separate and is explained in section 10 and the Cancellation, Refund and Cooling-off Policy. It is not restricted by the ordinary renewal-notice rule above.",
    )
    add_para(
        doc,
        "A youth Checkout creates one subscription covering every child listed in that Checkout. An ordinary or cooling-off cancellation ends the whole subscription and all listed children’s places; the online flow does not cancel one child separately. Contact us about any requested change to the listed children. No addition, removal, plan, price or discount change takes effect unless and until we confirm it in writing, including the effective date and resulting recurring total. We may require cancellation and a new Checkout. Stopping attendance or asking to remove one child is not by itself a contract change.",
    )

    doc.add_heading("10. Cooling-off and refunds", level=1)
    add_para(
        doc,
        "Where the Consumer Contracts (Information, Cancellation and Additional Charges) Regulations 2013 apply, a consumer generally may cancel until the end of 14 days after the day the online service contract is made. If the payer expressly requests that service begin on the service start date shown, even if that is before the cooling-off period ends, and then cancels within that period, we may deduct only a proportionate amount for services actually supplied, where the law allows. If that express request was not made, different refund consequences may apply. Any statutory refund will be made within 14 days where that is the legal deadline, to the original payment method unless the consumer expressly agrees otherwise, and without a refund fee.",
    )
    add_para(
        doc,
        "Outside a statutory or other mandatory right, payments are non-refundable. Nothing in these Terms excludes remedies for services not provided with reasonable care and skill, misdescription, breach of contract, or any other right that cannot lawfully be excluded.",
    )

    doc.add_heading("11. Failed payments, grace period and disputes", level=1)
    add_para(
        doc,
        "After a membership has started, if a recurring payment fails, Stripe may retry it and we may ask the payer to update the payment method. We allow a three-calendar-day past-due grace period from the failed due date, during which existing related access continues while payment is recovered. If the subscription remains past due after the grace period, access may be suspended until payment succeeds or the membership ends. If the first scheduled payment fails, the membership and AlphaWOD access do not start. We do not add an undisclosed late fee or accelerate future monthly payments.",
    )
    add_para(
        doc,
        "When a payment dispute is opened, related access is suspended while the dispute is investigated. If the dispute is resolved in our favour, access may be restored promptly, subject to current membership status, and we will fairly assess any credit or extension needed for paid access that was unavailable. If the dispute is lost or the payment is fully refunded, related access is revoked. This does not remove the payer’s right to raise a genuine complaint, use an applicable chargeback right, or exercise a statutory remedy. Contacting us first may allow a billing error to be resolved more quickly.",
    )

    doc.add_heading("12. Conduct, safety and use of facilities", level=1)
    add_bullets(
        doc,
        [
            "Participants must follow staff instructions, posted rules, equipment guidance and reasonable safeguarding measures.",
            "Participants must use equipment only as instructed, wear suitable clothing and footwear, and behave safely and respectfully.",
            "A participant should stop immediately and tell a member of staff if they feel unwell, unsafe or unable to continue.",
            "Violence, harassment, deliberate damage, dangerous conduct, unauthorised commercial activity or misuse of another person’s membership may result in proportionate restriction or termination after fair consideration of the circumstances.",
            "A guardian must comply with youth arrival, handover, collection and supervision arrangements communicated during onboarding.",
        ],
    )
    add_para(
        doc,
        "We may take immediate, proportionate action where reasonably necessary to protect a person, property, safeguarding or service security. Where appropriate, we will explain the decision and offer a review route.",
    )

    doc.add_heading("13. Health and participation", level=1)
    add_para(
        doc,
        "Physical training carries inherent risks. The participant or guardian is responsible for deciding whether to seek medical advice before participation and for following professional advice. The public purchase form does not request medical details. If information is necessary for safe participation, do not place it in a signature or general support field; use the separate secure onboarding route we provide. The applicable waiver or guardian addendum contains the detailed participation acknowledgement.",
    )

    doc.add_heading("14. Our responsibility", level=1)
    add_para(
        doc,
        "We must provide services with reasonable care and skill. Nothing in these Terms excludes or limits liability for death or personal injury caused by negligence, fraud or fraudulent misrepresentation, breach of statutory rights, or any liability that cannot lawfully be excluded or limited.",
    )
    add_para(
        doc,
        "Subject to that protection, we are not responsible for loss caused by the participant’s deliberate or unsafe misuse of facilities, breach of clear safety instructions, or an event outside our reasonable control where we took reasonable steps to reduce the effect. We are not responsible for business losses arising from a consumer membership. Any limitation will apply only so far as fair and lawful in the circumstances.",
    )

    doc.add_heading("15. Customer Portal and account security", level=1)
    add_para(
        doc,
        "The Stripe Customer Portal may allow the payer to update a payment method and view invoices. Pause and Stripe’s standard cancellation control are disabled because cancellation uses the notice process in section 9. The payer may open the Customer Portal only from the signed-in account that owns the membership. A purchaser who checked out while signed out must first claim the purchase from the Checkout return flow or from a signed-in account whose verified email matches the payer email. The recipient must keep account credentials secure and tell us promptly about suspected unauthorised use.",
    )

    doc.add_heading("16. Ending or restricting a membership by us", level=1)
    add_para(
        doc,
        "We may suspend or end a membership for non-payment, serious or repeated breach, fraud, misuse, safety or safeguarding risk, or where continuing the service is unlawful or no longer reasonably possible. Except where immediate action is reasonably necessary, we will give notice, explain the reason and allow a reasonable opportunity to put a remediable breach right. Any refund or continuing charge will be assessed fairly and in accordance with the law.",
    )

    doc.add_heading("17. Changes to these Terms", level=1)
    add_para(
        doc,
        "We may update these Terms for legal, regulatory, security or reasonable service reasons. We will not impose a material adverse change retrospectively. Where a change materially affects an active membership, we will give clear advance notice and any fair cancellation right required by law. The version accepted at checkout and any later version lawfully notified will be retained with the acceptance evidence.",
    )

    doc.add_heading("18. Communications, complaints and governing law", level=1)
    add_para(
        doc,
        f"We send service, billing and legal communications to the payer’s recorded email. The payer must keep it current. Complaints should be sent to {SUPPORT_EMAIL} or by post to {ADDRESS}. We will investigate and respond within a reasonable time.",
    )
    add_para(
        doc,
        "These Terms are governed by the law of England and Wales. A consumer living elsewhere in the UK retains any mandatory protection of the part of the UK where they live and may bring proceedings in any court available under applicable consumer law. Nothing in these Terms prevents either party from using another lawful dispute-resolution route.",
    )

    add_acceptance_block(
        doc,
        "Checkout acceptance text",
        [
            "I have read and agree to the Membership Terms and the Cancellation, Refund and Cooling-off Policy. I confirm that all participant and payer or guardian details I supplied are accurate.",
            "I acknowledge that I have received and read the Privacy Notice explaining how personal information is used.",
            "I authorise the amount Stripe shows today and future recurring monthly payments for the selected membership on the billing schedule shown at Checkout. For a youth membership, Stripe will show the plan, quantity, price per child, undiscounted subtotal, recurring 15% discount when the quantity is two or more, and resulting recurring total. For another verified promotion, Stripe will show its effect and when the standard price resumes. This authority is subject to my cancellation and statutory rights.",
            "I expressly request that the membership and any eligible AlphaWOD access begin on the service start date shown, even if that is before the 14-day cooling-off period ends. I understand that, if I cancel during that period, Zero Alpha Fitness may retain or charge only the proportionate amount permitted by law for services supplied before cancellation.",
        ],
        "payer name and verified payer email",
    )

    add_review_appendix(
        doc,
        [
            "Obtain counsel’s fairness assessment of the 14-day pre-renewal cutoff and the extra monthly charge where notice arrives later, including prominence at catalogue, checkout, confirmation and cancellation request.",
            "Confirm the operational definition of receipt, deadline timestamps, timezone and what happens if the cancellation service or email system is unavailable.",
            "Confirm each membership’s facilities, class access, capacity rules and any required booking/no-show terms before publication.",
            "Set the advance-notice period and customer remedy for future price or material service changes.",
            "Confirm that adult self-purchase is enforced consistently across Checkout, account claim, support and any assisted route.",
            "Set an explicit rule for a participant aged 17. The approved youth bands end at 16 and this draft requires an adult participant to be 18, so age 17 is intentionally not offered until the owner and counsel approve a route.",
            "Confirm that the registered office, place of registration and public trading/contact address remain current on the website and commercial documents immediately before publication.",
            "Review the suspension/termination and liability clauses against current insurance, house rules, safeguarding policy and complaint procedure.",
            "Verify Welsh/English language, accessibility and durable-medium requirements for the intended customer base and checkout design.",
            "Review these Terms again against any commenced subscription-contract provisions of the Digital Markets, Competition and Consumers Act 2024 before launch and on each material update.",
        ],
        [
            ("Consumer Contracts Regulations 2013", "https://www.legislation.gov.uk/uksi/2013/3134/contents"),
            ("Consumer Rights Act 2015", "https://www.legislation.gov.uk/ukpga/2015/15/contents"),
            ("CMA unfair contract terms guidance", "https://www.gov.uk/government/publications/unfair-contract-terms-cma37"),
            ("CMA fair-contract guidance updated 22 July 2026", "https://www.gov.uk/guidance/writing-a-fair-contract-for-customers"),
            ("Unfair Contract Terms Act 1977", "https://www.legislation.gov.uk/ukpga/1977/50/contents"),
            ("Digital Markets, Competition and Consumers Act 2024", "https://www.legislation.gov.uk/ukpga/2024/13/contents"),
            ("Government response on the subscription-contract regime", "https://www.gov.uk/government/consultations/consultation-on-the-implementation-of-the-new-subscription-contracts-regime/outcome/government-response-to-consultation-on-the-implementation-of-the-new-subscription-contracts-regime-web-accessible-version"),
            ("GOV.UK online-selling requirements", "https://www.gov.uk/online-and-distance-selling-for-businesses/online-selling"),
            ("Stripe subscription billing-cycle documentation", "https://docs.stripe.com/billing/subscriptions/billing-cycle"),
        ],
        high_priority=(
            "The owner-proposed policy is a 14-day deadline before the next first-of-month renewal, not cancellation that always takes effect exactly 14 days after request. A late request can therefore keep the contract running for one extra full billing month. The customer-facing wording makes that effect prominent, but counsel must assess fairness and current or forthcoming subscription-contract rules before it is used."
        ),
    )
    path = OUTPUT_DIR / spec.filename
    doc.save(path)
    return path


def build_privacy() -> Path:
    spec = SPECS["privacy"]
    doc = Document()
    configure_document(doc, spec)
    # The privacy notice is the densest draft. Slightly tighter, still-readable
    # body rhythm keeps the short final clause with the customer-facing notice
    # instead of orphaning it on an otherwise blank page before the appendix.
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.25)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.07
    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.size = Pt(10.25)
        style.paragraph_format.space_after = Pt(3.5)
        style.paragraph_format.line_spacing = 1.05
    add_document_title(doc, spec)
    add_draft_banner(doc)

    doc.add_heading("1. Who we are", level=1)
    add_para(
        doc,
        f"{ENTITY}, a company registered in England and Wales under company number {COMPANY_NUMBER}, trading as {TRADING_NAME}, is the controller of the personal information described in this Notice. Our registered office is {REGISTERED_OFFICE}. Our public trading and contact address is {ADDRESS}. For privacy questions or requests, email {SUPPORT_EMAIL}.",
    )
    add_para(
        doc,
        "This Notice covers membership buyers and payers, adult participants, children in Youngstars or Teenstars, their parents or guardians, AlphaWOD account holders, and people who contact support or exercise a privacy right.",
    )

    doc.add_heading("Checkout short-form notice", level=2)
    add_quote(
        doc,
        "ZERO ALPHA FITNESS LTD uses these details to set up and manage the named participant’s or participants’ membership, verify age and guardian authority for each child, manage payment through Stripe, and provide eligible AlphaWOD access. Required fields are needed to complete the membership. If you provide details about another person, we will give them our Privacy Notice during onboarding. Marketing is not part of this purchase. Read the Privacy Notice.",
    )

    doc.add_heading("2. The information we collect", level=1)
    doc.add_heading("Identity, eligibility and contact information", level=2)
    add_bullets(
        doc,
        [
            "Each participant’s name and date of birth, and the payer email collected through Stripe or the account-claim flow.",
            "For youth membership, each child’s selected plan and age eligibility, the guardian’s name and relationship to each child, and the guardian’s declaration of authority for every named child.",
            "The Zero Alpha purchase form does not request a phone number or billing address. Stripe may ask the payer directly for information required by a selected payment method or its legal and fraud-prevention checks.",
        ],
    )
    doc.add_heading("Account, membership and agreement information", level=2)
    add_bullets(
        doc,
        [
            "Firebase account identifier, verified email status, sign-in and security records. Authentication providers handle credentials; we do not need to see the payer’s full password.",
            "Selected plan, every named participant, contracted quantity, price per child, undiscounted subtotal, multi-child-discount eligibility, rate and amount, resulting recurring total, start date, billing anchor, other promotion use, subscription and access status, payment-failure grace period, cancellation request and effective end date.",
            "Typed electronic signature, authority declaration, acceptance statements, accepted document version, timestamp and the audit context actually recorded at acceptance.",
            "AlphaWOD eligibility, account-claim state, bookings, attendance, workout or training entries and performance records for eligible adult users.",
        ],
    )
    doc.add_heading("Payment information", level=2)
    add_para(
        doc,
        "Stripe collects and processes payment details. We expect to receive identifiers and status information such as Stripe customer, Checkout Session, subscription, invoice, payment, refund and dispute references, and limited payment-method information where Stripe makes it available. We do not store full card or bank credentials in the Zero Alpha application.",
    )
    doc.add_heading("Technical, communications and safety information", level=2)
    add_bullets(
        doc,
        [
            "IP address, device/browser information, timestamps, authentication events, security, error and audit logs to the extent the service providers and our implementation record them.",
            "Transactional email delivery events, support correspondence, complaints and privacy requests.",
            "The public purchase form does not request health or injury details. If safety information is needed during onboarding, it must be collected through the separate, clearly explained process provided for that purpose.",
            "Photographs, video or promotional media only through a separate optional consent process; media consent is not a condition of membership.",
        ],
    )

    doc.add_heading("3. Where the information comes from", level=1)
    add_bullets(
        doc,
        [
            "Directly from an adult participant buying for themselves, or from a guardian buying for one or more children, during purchase, onboarding, account use or support contact.",
            "From Stripe about checkout, billing, payment, refund and dispute events.",
            "Automatically from the site, AlphaWOD, authentication service and security or operational logs.",
            "From authorised staff recording membership, onboarding, attendance, booking or training administration.",
        ],
    )
    add_para(
        doc,
        "Adult participants receive this Notice during their own Checkout. For children, the guardian receives the full Notice and we will provide age-appropriate information to each child during onboarding.",
    )

    doc.add_heading("4. Why we use information and our lawful bases", level=1)
    add_para(doc, "We use personal information only where a lawful basis applies:")
    add_bullets(
        doc,
        [
            "Contract — to take the payer’s payment, create and administer their subscription, send essential billing/service messages, process cancellation and deliver services that are objectively necessary under a contract with that person.",
            "Legitimate interests — to administer a youth membership, confirm guardian authority and age eligibility, prevent duplicate subscriptions and fraud, secure the service, keep proportionate audit evidence, administer ordinary bookings and training, manage claims and improve reliable operations. Our interests are providing the agreed service, protecting users and the business, and demonstrating what was agreed. We must balance those interests against each person’s rights, with extra weight for children.",
            "Legal obligation — to keep company, accounting and transaction records and make disclosures where a specific law requires it.",
            "Consent — only where a genuinely optional activity requires it, such as optional marketing, promotional media, or a separately designed optional feature using health information. Consent can be withdrawn without affecting earlier lawful use.",
            "Vital interests — exceptionally, where processing is necessary to protect someone’s life and they cannot consent, for example in a genuine emergency.",
        ],
    )
    add_para(
        doc,
        "Service, security, payment and cancellation messages are not marketing. If we introduce marketing, we will keep it separate from purchase acceptance and use consent or another route only where PECR and data-protection law allow it. Every marketing message will offer the required opt-out.",
    )

    doc.add_heading("5. Children and youth membership", level=1)
    add_para(
        doc,
        "Youngstars covers ages 6–11 and Teenstars covers ages 12–16. One youth Checkout may contain one or more children in one selected plan only; Youngstars and Teenstars are not combined in one Checkout. The guardian is the payer and supplies each child’s details. Children do not receive AlphaWOD access under the initial public youth membership. We use only the information reasonably needed to verify each child’s eligibility, arrange onboarding, provide the membership, protect each child and maintain appropriate evidence.",
    )
    add_para(
        doc,
        "A child’s data-protection rights belong to the child. Whether a guardian may exercise a right for them depends on the child’s understanding, authority and best interests. We will explain relevant processing in language and a format appropriate to the child’s age. We do not use a guardian’s acceptance of the membership documents as blanket data-protection consent.",
    )

    doc.add_heading("6. Health and other special-category information", level=1)
    add_callout(
        doc,
        "Do not submit health details at checkout",
        "Do not type medical conditions, injuries, medication or other health information into a signature, promotion-code, support or general checkout field. If information is needed for safe participation, use the separate secure onboarding route provided by Zero Alpha Fitness.",
        kind="warning",
    )
    add_para(
        doc,
        "Information about health is special-category data. If we collect it, we need both an Article 6 lawful basis and an Article 9 condition, plus clear information and appropriate safeguards. Some workout or performance entries may reveal health even if they are not labelled medical. Before any such feature is used, we will classify the fields, identify the lawful condition, minimise access and retention, and provide a just-in-time notice. Where explicit consent is the appropriate condition for an optional feature, it will be specific, separate and withdrawable.",
    )

    doc.add_heading("7. Automated status changes and human review", level=1)
    add_para(
        doc,
        "Stripe events and our rules may automatically update subscription status and related access—for example, allowing Adult Unlimited access after confirmed payment, applying a three-day past-due grace period, blocking a duplicate active subscription, suspending access when a dispute opens, restoring it after a dispute is won, or revoking it after a lost dispute or full refund. These rules use payment and membership status rather than profiling a person’s character. A person may contact us for an explanation, correction and human review of an incorrect or exceptional result.",
    )

    doc.add_heading("8. Who receives information", level=1)
    add_para(doc, "We share only what is reasonably necessary with:")
    add_bullets(
        doc,
        [
            "Stripe, payment networks, banks and payment-method providers for Checkout, Billing, payment authentication, fraud prevention, refunds, disputes and the Customer Portal. Stripe may act as our service provider and as an independent controller for some legal, network and fraud purposes under its own notice.",
            "Google Firebase and Google Cloud for authentication, database, server functions and related infrastructure.",
            "Vercel for website delivery, hosting and operational logs.",
            "Resend for transactional email delivery.",
            "Authorised Zero Alpha Fitness staff and contractors who need the information for their role.",
            "Accountants, insurers, legal advisers, courts, regulators, law-enforcement bodies or a buyer/reorganised business where disclosure is lawful, necessary and appropriately protected.",
        ],
    )
    add_para(
        doc,
        "A supplier is not automatically our processor for every activity. We verify each role, contract and subprocessor and keep access limited to the relevant purpose.",
    )

    doc.add_heading("9. International transfers", level=1)
    add_para(
        doc,
        "Some suppliers or their subprocessors may process information outside the UK. Where this involves a restricted transfer, we use an applicable UK adequacy regulation or an approved UK International Data Transfer Agreement or UK Addendum, together with any required risk assessment and supplementary protections. A person may ask us for information about the relevant safeguard.",
    )

    doc.add_heading("10. How long we keep information", level=1)
    add_para(
        doc,
        "We keep information only for as long as needed for the purpose, legal records, security and live disputes, then delete or irreversibly anonymise it. Our retention schedule is:",
    )
    add_bullets(
        doc,
        [
            "Company, accounting, invoice and transaction records — normally six years from the end of the company financial year to which they relate, subject to any longer lawful requirement or live enquiry.",
            "Membership account, subscription, cancellation and essential support records — while active and normally up to six years after the membership or account relationship ends where needed for contract, complaint or claim records; unnecessary live-profile data should be removed sooner.",
            "Adult waiver, guardian authority, acceptance version and signature evidence — six years after adult membership ends; for a child, until at least the 21st birthday or resolution of a live claim, whichever is later, subject to any longer lawful, insurance or safeguarding requirement.",
            "Bookings, attendance and ordinary workout/performance entries — for the operational period communicated in AlphaWOD settings and then deleted or anonymised.",
            "Authentication, security and technical logs — a short, documented period appropriate to the risk, normally no more than 12 months unless needed for an incident, fraud or claim.",
            "Optional marketing — until consent is withdrawn or the purpose ends, while retaining only a minimal suppression record where needed to honour an opt-out.",
        ],
    )
    add_para(
        doc,
        "A deletion request does not override a lawful need to retain a limited record. Where possible, financial and claims evidence will be separated from the live app profile so unnecessary operational information can be removed.",
    )

    doc.add_heading("11. Cookies, local storage and similar technologies", level=1)
    add_para(
        doc,
        "The site, Firebase authentication, Stripe and hosting services may use cookies, browser storage, device identifiers or similar storage and access technologies for requested sessions, security, payment authentication, fraud prevention and reliable delivery. Strictly necessary technologies do not require consent but still require clear information. Any non-exempt analytics, advertising or cross-service tracking remains disabled unless and until valid consent is obtained.",
    )

    doc.add_heading("12. Security", level=1)
    add_para(
        doc,
        "We use proportionate technical and organisational measures designed to protect information, including access controls, verified sign-in and one-time purchase-claim verification, separation of payment credentials from the app, audit logging and supplier security terms. No online service is risk-free. If a personal-data breach creates a legally reportable risk, we will notify the ICO and affected people as required.",
    )

    doc.add_heading("13. Your rights", level=1)
    add_para(doc, "Depending on the circumstances, a person may have the right to:")
    add_bullets(
        doc,
        [
            "be informed and obtain access to their personal information;",
            "correct inaccurate or incomplete information;",
            "ask for erasure or restriction;",
            "receive portable information where the legal conditions apply;",
            "withdraw consent at any time where consent is used;",
            "object to direct marketing; and",
            "ask for safeguards and human review concerning qualifying automated decisions.",
        ],
    )
    add_callout(
        doc,
        "Right to object",
        f"You may object to processing based on our legitimate interests. Tell us what you object to and why by emailing {SUPPORT_EMAIL}. We will stop unless we demonstrate compelling legitimate grounds that override your interests, rights and freedoms, or the processing is needed for legal claims. Direct marketing stops when you object.",
        kind="info",
    )
    add_para(
        doc,
        "We may need reasonable information to verify identity and authority before acting. We do not charge for an ordinary request, but the law allows limited exceptions. We will respond within the applicable time and explain any lawful refusal or extension.",
    )

    doc.add_heading("14. Complaints", level=1)
    add_para(
        doc,
        f"Send a data-protection complaint to {SUPPORT_EMAIL} or {ADDRESS}. We will provide a direct route, acknowledge the complaint within 30 days, investigate appropriately, keep the complainant informed and provide the outcome without undue delay. A person may also complain to the UK Information Commissioner’s Office (ICO).",
    )
    p = doc.add_paragraph()
    p.add_run("ICO complaint information: ")
    add_hyperlink(
        p,
        "ico.org.uk/make-a-complaint/data-protection-complaints/",
        "https://ico.org.uk/make-a-complaint/data-protection-complaints/",
    )

    doc.add_heading("15. Changes to this Notice", level=1)
    add_para(
        doc,
        "We will version and date this Notice. If a change materially affects how active-member information is used, we will provide a clear notice before the new use where required. Earlier acceptance and transaction records remain linked to the version provided at the relevant time.",
    )

    add_review_appendix(
        doc,
        [
            "Confirm support@zeroalphafitness.co.uk as the privacy contact, nominate the internal owner, and confirm whether a data protection officer is required or appointed.",
            "Complete the data inventory: exact checkout, Firebase, Stripe, AlphaWOD, log, email, IP/user-agent and payment-method metadata fields actually stored.",
            "Complete a child-focused legitimate-interests assessment and DPIA; prepare distinct child-friendly explanations for ages 6–11 and 12–16 before youth onboarding launches.",
            "Classify AlphaWOD workout/performance fields and design an Article 9 process before collecting any data that reveals health.",
            "Verify each supplier’s current legal entity, role, data-processing agreement, subprocessor list, hosting/log regions and UK transfer mechanism.",
            "Complete the cookie/device-storage inventory and implement consent before any non-exempt technology operates.",
            "Approve and technically enforce the retention/deletion schedule, including backups, abandoned checkouts, support records, child signature evidence and account deletion.",
            "Assess whether Stripe or internal access rules produce solely automated decisions with legal or similarly significant effects, and ensure meaningful human review.",
            "Confirm whether CCTV, incident reports, emergency contacts, photographs/video, marketing or other channels require separate notices or policies.",
            "Verify and document the direct privacy-complaint process introduced in 2026 and the customer-facing ICO details immediately before publication.",
        ],
        [
            ("ICO right to be informed", "https://ico.org.uk/for-organisations/uk-gdpr-guidance-and-resources/individual-rights/individual-rights/right-to-be-informed/"),
            ("ICO lawful basis: contract", "https://ico.org.uk/for-organisations/uk-gdpr-guidance-and-resources/lawful-basis/a-guide-to-lawful-basis/contract/"),
            ("ICO children and the UK GDPR", "https://ico.org.uk/for-organisations/uk-gdpr-guidance-and-resources/childrens-information/children-and-the-uk-gdpr/"),
            ("ICO special-category conditions", "https://ico.org.uk/for-organisations/uk-gdpr-guidance-and-resources/lawful-basis/special-category-data/what-are-the-conditions-for-processing/"),
            ("ICO international transfers", "https://ico.org.uk/for-organisations/uk-gdpr-guidance-and-resources/international-transfers/"),
            ("ICO storage limitation", "https://ico.org.uk/for-organisations/uk-gdpr-guidance-and-resources/data-protection-principles/a-guide-to-the-data-protection-principles/storage-limitation/"),
            ("ICO storage and access technologies", "https://ico.org.uk/for-organisations/direct-marketing-and-privacy-and-electronic-communications/guidance-on-the-use-of-storage-and-access-technologies/"),
            ("ICO data-protection complaints", "https://ico.org.uk/for-organisations/how-to-deal-with-data-protection-complaints/"),
            ("GOV.UK company and accounting records", "https://www.gov.uk/running-a-limited-company/company-and-accounting-records"),
        ],
        high_priority=(
            "The data map, international-transfer audit, exact retention schedule, cookies/device-storage inventory and child-specific DPIA are not yet confirmed. This draft deliberately labels those gaps; publish only after the text matches the implemented systems and current supplier contracts."
        ),
    )
    path = OUTPUT_DIR / spec.filename
    doc.save(path)
    return path


def build_cancellation() -> Path:
    spec = SPECS["cancel"]
    doc = Document()
    configure_document(doc, spec)
    # Keep the short final contact section with the substantive policy instead
    # of allowing Word/LibreOffice pagination to leave it alone on a fourth page.
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.25)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.07
    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.size = Pt(10.25)
        style.paragraph_format.space_after = Pt(3.5)
        style.paragraph_format.line_spacing = 1.05
    add_document_title(doc, spec)
    add_draft_banner(doc)

    doc.add_heading("At a glance", level=1)
    add_bullets(
        doc,
        [
            "Memberships are rolling monthly, with no joining fee, free trial, minimum term or pause option.",
            "During the founding presale, £0 is charged today; membership starts and the first monthly payment is taken on 1 September 2026. After opening, a person joining after the first pays the Stripe-displayed prorated amount immediately.",
            "To avoid the next first-of-month payment, the cancellation request must reach us at least 14 calendar days before that billing date.",
            "If the request arrives later, the next first-of-month payment remains due and membership continues through that additional paid month.",
            "Payments are non-refundable except where required by law.",
            "A statutory 14-day cooling-off right for a new online contract is separate from the ordinary renewal-notice rule.",
            "A youth family Checkout is one subscription. Ordinary or cooling-off cancellation ends every listed child’s place; the online flow does not cancel one child separately.",
        ],
    )

    doc.add_heading("1. Presale, start date, proration and monthly renewal", level=1)
    add_para(
        doc,
        "During the founding presale, membership and service start on 1 September 2026 and the first monthly payment is taken then. A cancellation request received before that service-start date withdraws the presale: no first payment is taken and service does not begin. After opening, membership begins when Checkout is confirmed; Stripe calculates any immediate proration until the next first. The full monthly price is then collected on each following first while the membership continues. Times and deadlines use Europe/London.",
    )

    doc.add_heading("2. How to request ordinary cancellation", level=1)
    add_numbered(
        doc,
        [
            "Use the signed-in cancellation-request flow. If you cannot access it, email support@zeroalphafitness.co.uk from the payer’s recorded email.",
            "Identify the payer, every participant and the membership. Do not send card or bank details.",
            "The signed-in flow records the server receipt time and freezes the displayed outcome. For email, staff record the time the message reaches the support inbox, not when it was written or sent from a device.",
            "The online flow sends an acknowledgement automatically. Staff send the email-channel acknowledgement after intake. It shows the recorded receipt time, the final scheduled payment (if any) and the membership end date. Contact us promptly if it is wrong.",
        ],
    )
    add_para(
        doc,
        "Stripe’s Customer Portal allows payment-method updates and invoice viewing. Its built-in pause and cancellation controls are disabled. Opening the portal or removing a payment method is not a cancellation request.",
    )
    add_para(
        doc,
        "A youth Checkout creates one subscription covering all children listed in that Checkout. An ordinary cancellation ends that whole subscription and all listed children’s places. The online flow does not cancel one child separately. Stopping one child’s attendance or asking us to remove one child does not itself change the contract, quantity, total or discount. Contact us about a requested roster change; it takes effect only when we confirm the effective date and resulting recurring total in writing, and we may require cancellation and a new Checkout.",
    )

    doc.add_heading("3. The 14-day renewal deadline", level=1)
    add_callout(
        doc,
        "Important financial commitment",
        "A request must reach us at least 14 calendar days before the next first of the month to stop that payment. If it reaches us less than 14 days before the next first, that payment remains due and cancellation takes effect after the additional paid month.",
        kind="warning",
    )
    add_bullets(
        doc,
        [
            "On-time request: no payment on the next first; access ends at 23:59 Europe/London on the day before it.",
            "Late request: the next first-of-month payment is collected; access continues for that full paid month and ends at 23:59 Europe/London on the day before the following first.",
            "After acknowledgement, no payment should be scheduled beyond the stated final payment. A billing error will be corrected, including any refund required by law or by our confirmation.",
        ],
    )
    doc.add_heading("Worked examples", level=2)
    add_bullets(
        doc,
        [
            "Next billing date 1 June; request received 18 May: this is 14 calendar days before 1 June. No 1 June payment is due and access ends 31 May.",
            "Next billing date 1 June; request received 19 May: this is less than 14 calendar days before 1 June. The 1 June payment remains due and access ends 30 June.",
        ],
    )
    add_para(
        doc,
        "The account flow should display the exact deadline and outcome before submission. If the service was unavailable, send the request by email and retain evidence of delivery; we will assess the circumstances fairly.",
    )

    doc.add_heading("4. Statutory cooling-off for a new online membership", level=1)
    add_para(
        doc,
        "If the Consumer Contracts (Information, Cancellation and Additional Charges) Regulations 2013 apply, a consumer generally may cancel an online service contract during the 14-day period beginning after the contract is made. This is a separate right. The ordinary renewal deadline in section 3 does not shorten it.",
    )
    add_para(
        doc,
        "Checkout asks the payer separately to request that membership and eligible AlphaWOD access begin on the service start date shown, even if that is before the cooling-off period ends. If the payer makes that express request and cancels during the cooling-off period, we may keep or charge only the proportionate amount the law permits for services actually supplied before cancellation. For a youth family subscription, cooling-off cancellation ends the whole subscription and every listed child’s place. A cooling-off request records an immediate cancellation outcome, but any refund or deduction is reviewed manually. The balance will be refunded within 14 days, to the original payment method unless the consumer expressly agrees otherwise, and without a refund fee. If the express service-start request was not made, we will not make a deduction that the Regulations prohibit.",
    )
    add_para(
        doc,
        "A service consumer loses the cooling-off right because of full performance within the 14 days only where all legal conditions are met, including the required express request and acknowledgement. Nothing in this policy asks a customer to give up a right unlawfully.",
    )
    add_quote(
        doc,
        "☐ I expressly request that the membership and any eligible AlphaWOD access begin on the service start date shown, even if that is before the 14-day cooling-off period ends. I understand that, if I cancel during that period, Zero Alpha Fitness may retain or charge only the proportionate amount permitted by law for services supplied before cancellation.",
    )

    doc.add_heading("5. How to use the cooling-off right", level=1)
    add_para(
        doc,
        "Make a clear statement that you want to cancel during the cooling-off period through the cancellation flow or by emailing support@zeroalphafitness.co.uk. You may use the model wording below, but you do not have to use it:",
    )
    add_quote(
        doc,
        "To ZERO ALPHA FITNESS LTD: I give notice that I cancel my membership contract. Payer name: [name]. Participant name(s): [all names]. Contract date: [date]. Payer email: [email]. Date of notice: [date].",
    )

    doc.add_heading("6. Refunds", level=1)
    add_para(
        doc,
        "Payments are non-refundable except where required by law. This means we do not ordinarily refund a correctly calculated after-opening proration, a used or unused part of a paid month, a promotion difference, or a payment that remained due because a cancellation request missed the 14-day renewal deadline. A valid pre-start presale withdrawal has no first payment to refund.",
    )
    add_para(
        doc,
        "We will provide any remedy required for a valid cooling-off cancellation, duplicate or incorrect charge, failure to provide services with reasonable care and skill, material breach, or other statutory right. A contractual no-refund statement never overrides a mandatory remedy.",
    )

    doc.add_heading("7. No pauses", level=1)
    add_para(
        doc,
        "Membership cannot be paused, frozen or placed on holiday hold. A request to pause does not stop billing and is not treated as a cancellation unless it clearly asks us to cancel. If disability, pregnancy, illness or another exceptional circumstance engages a legal duty or makes the standard policy unfair, contact us so we can consider a reasonable and lawful response.",
    )

    doc.add_heading("8. Failed payments and disputes", level=1)
    add_para(
        doc,
        "After a membership has started, a failed recurring payment enters a three-calendar-day past-due grace period, during which existing related access continues. Stripe may retry payment and we may ask the payer to update the method. If payment is still past due after the grace period, access may be suspended. If the first scheduled payment fails, the membership and AlphaWOD access do not start. An open payment dispute suspends related access; a dispute won by Zero Alpha Fitness restores eligible access promptly and we will fairly assess any credit or extension needed for paid time that was unavailable; a lost dispute or full refund revokes related access. These access rules do not prevent a genuine complaint, statutory cancellation or lawful chargeback.",
    )

    doc.add_heading("9. Confirmation and contact", level=1)
    add_para(
        doc,
        f"Cancellation and refund communications will be sent to the payer’s recorded email. Questions or complaints: {SUPPORT_EMAIL}. Post: {ENTITY}, {ADDRESS}. Keep the acknowledgement and relevant Stripe receipt.",
    )

    add_review_appendix(
        doc,
        [
            "Counsel must assess the fairness and enforceability of collecting one additional full month where notice arrives less than 14 days before renewal, including whether a less burdensome rule should replace it.",
            "Confirm the exact deadline calculation, treatment of leap years, DST, delivery failures and support inbox outages; the app must compute and display the same result server-side.",
            "Confirm the precise contract-formation timestamp and durable-medium delivery of cancellation instructions and the model form.",
            "Confirm the immediate-performance control is separate, unticked and evidenced, and that cooling-off refund calculations use services actually supplied rather than an arbitrary fixed fee.",
            "Confirm refund timing and payment method rules in the implementation and customer support playbook.",
            "Review exceptional-circumstance handling against equality duties, unfair-terms law and insurance/operational policy.",
            "Re-review before launch and before spring 2027 against the commencement and transition provisions for subscription contracts under the Digital Markets, Competition and Consumers Act 2024.",
        ],
        [
            ("Consumer Contracts Regulations 2013", "https://www.legislation.gov.uk/uksi/2013/3134/contents"),
            ("Government guidance on the Consumer Contracts Regulations", "https://www.gov.uk/government/publications/consumer-contracts-regulations-2013"),
            ("Consumer Rights Act 2015", "https://www.legislation.gov.uk/ukpga/2015/15/contents"),
            ("CMA unfair contract terms guidance", "https://www.gov.uk/government/publications/unfair-contract-terms-cma37"),
            ("CMA fair-contract guidance updated 22 July 2026", "https://www.gov.uk/guidance/writing-a-fair-contract-for-customers"),
            ("Digital Markets, Competition and Consumers Act 2024", "https://www.legislation.gov.uk/ukpga/2024/13/contents"),
            ("Government response on the subscription-contract regime", "https://www.gov.uk/government/consultations/consultation-on-the-implementation-of-the-new-subscription-contracts-regime/outcome/government-response-to-consultation-on-the-implementation-of-the-new-subscription-contracts-regime-web-accessible-version"),
            ("Stripe cancellation documentation", "https://docs.stripe.com/billing/subscriptions/cancel"),
        ],
        high_priority=(
            "The owner-proposed late-notice outcome is materially different from cancellation taking effect exactly 14 days after a request: it can continue the contract through the following monthly cycle. The draft uses that billing outcome and makes it prominent, but it should not be published without a current UK consumer-law fairness review."
        ),
    )
    path = OUTPUT_DIR / spec.filename
    doc.save(path)
    return path


def build_adult_waiver() -> Path:
    spec = SPECS["adult"]
    doc = Document()
    configure_document(doc, spec)
    add_document_title(doc, spec)
    add_draft_banner(doc)

    add_callout(
        doc,
        "Participant must accept personally",
        "Every adult membership must be bought and signed by the named adult participant for themselves. Third-party purchase of an adult membership is not supported.",
        kind="info",
    )

    doc.add_heading("1. Participant declaration", level=1)
    add_para(
        doc,
        "I confirm that I am the named participant, I am aged 18 or over, and the information I provide is accurate. I understand that this document records informed participation and reasonable allocation of risk; it does not remove duties or rights that the law does not allow either party to exclude.",
    )

    doc.add_heading("2. Activities covered", level=1)
    add_para(
        doc,
        "This acknowledgement applies to the Zero Alpha Fitness activities I choose to undertake under my membership, which may include gym use, resistance and cardiovascular training, functional fitness, coached classes, individual or group workouts, use of free weights and machines, conditioning, mobility work and related warm-up or cool-down activity, whether at the facility or at an organised session.",
    )

    doc.add_heading("3. Risks I understand", level=1)
    add_para(
        doc,
        "Physical training has inherent risks even where reasonable care is taken. Depending on the activity, these may include slips, trips, falls, collisions, equipment movement or failure, overexertion, delayed-onset soreness, strains, sprains, fractures, head or spinal injury, aggravation of an existing condition, heat illness, cardiovascular events and, in rare cases, permanent injury or death. Risks may arise from my own actions, other participants, the environment and the nature of strenuous exercise.",
    )
    add_para(
        doc,
        "I voluntarily choose to participate with that understanding. I accept inherent risks that cannot be eliminated by reasonable care, but I do not waive liability for negligence or any statutory right that cannot lawfully be excluded.",
    )

    doc.add_heading("4. Fitness to participate and medical advice", level=1)
    add_bullets(
        doc,
        [
            "I will consider my current fitness, health, experience and the demands of an activity before taking part.",
            "I will seek medical advice before participation where I have symptoms, concerns, a relevant condition, am pregnant or have been advised to limit exercise.",
            "I will follow medical advice and will not participate while impaired by alcohol, non-prescribed drugs, unsafe medication effects, acute illness or injury.",
            "I will tell the appropriate member of staff, through the secure route provided, about information reasonably necessary for safe participation and any change that affects it.",
        ],
    )
    add_callout(
        doc,
        "Health-information channel",
        "The public checkout and typed-signature fields are not designed for medical details. I will not enter health information there. If staff need safety information, I will use the separate onboarding channel and read the just-in-time privacy information.",
        kind="warning",
    )

    doc.add_heading("5. My conduct and safety responsibilities", level=1)
    add_bullets(
        doc,
        [
            "I will follow reasonable instructions, posted rules, equipment guidance and any scaling or exclusion communicated for safety.",
            "I will use equipment only for its intended purpose and only when I know how to do so or have asked for instruction.",
            "I will wear suitable clothing and footwear, keep the training area reasonably clear and act with care toward others.",
            "I will stop immediately and tell staff if I feel pain, dizziness, unusual shortness of breath, faintness, loss of control or any other warning sign.",
            "I will not deliberately conceal a safety issue, attempt a movement beyond my safe capability after being told not to, or disrupt another participant’s safe use of the facility.",
        ],
    )

    doc.add_heading("6. Coaching and results", level=1)
    add_para(
        doc,
        "Coaching and programming are general fitness services, not medical diagnosis or treatment. Results vary and are not guaranteed. I remain responsible for choosing weights, intensity and participation within my capability while following coaching and safety instructions.",
    )

    doc.add_heading("7. Emergency assistance and first aid", level=1)
    add_para(
        doc,
        "If staff reasonably believe urgent assistance is needed, I authorise them to provide or arrange proportionate first aid, contact emergency services and share the minimum information reasonably necessary for that emergency. This does not guarantee the availability of a particular treatment or professional and does not replace my responsibility to obtain medical advice.",
    )

    doc.add_heading("8. Personal property", level=1)
    add_para(
        doc,
        "I remain responsible for personal property I bring to the facility. Zero Alpha Fitness will take reasonable care where it assumes responsibility, but is not responsible for unattended loss or damage caused without its breach of duty. Nothing in this paragraph excludes liability that cannot lawfully be excluded.",
    )

    doc.add_heading("9. Liability and statutory rights", level=1)
    add_para(
        doc,
        "Zero Alpha Fitness must provide its services with reasonable care and skill. Nothing in this document excludes or limits liability for death or personal injury caused by negligence, fraud or fraudulent misrepresentation, breach of statutory rights, or another liability that cannot be limited by law.",
    )
    add_para(
        doc,
        "So far as fair and lawful, Zero Alpha Fitness is not responsible for harm or loss caused by my deliberate or unsafe misuse, material breach of clear safety instructions, inaccurate information I knowingly provide, or an inherent risk that remained despite reasonable care. Any responsibility will be assessed according to the facts and applicable law, including each party’s contribution.",
    )

    doc.add_heading("10. Privacy and media", level=1)
    add_para(
        doc,
        "The Privacy Notice explains how participant, signature, account, training and incident information is used. This waiver is not consent to marketing, photography, video or promotional use. Any media permission must be optional, specific and separate, and refusing it does not affect membership.",
    )

    doc.add_heading("11. Electronic signature and continuing effect", level=1)
    add_para(
        doc,
        "I intend my typed name and affirmative submission to authenticate and sign this document electronically. I will have an opportunity to review the document before signing and receive or access a durable copy afterwards. The acceptance record will include the document version and timestamp.",
    )
    add_para(
        doc,
        "This acknowledgement continues while I participate under the membership, but it does not silently authorise a materially different risk or remove the need to notify me of a material document change. If one provision is unenforceable, the remaining provisions continue so far as lawful. The law and jurisdiction wording in the Membership Terms applies, subject to mandatory consumer rights.",
    )

    doc.add_heading("Electronic signing arrangement", level=1)
    add_quote(doc, "☐ I confirm that I am the named participant, I am aged 18 or over, and I have read and understood the Adult Participant Waiver and Risk Acknowledgement. I understand the activities and inherent risks and choose to participate, subject to my statutory rights and Zero Alpha Fitness’s duty to take reasonable care.")
    add_para(
        doc,
        "Required fields: participant full legal name and date of birth; typed signature name; the exact acceptance statement; signature date/time; membership/order reference; waiver document ID and version. The adult participant is also the payer.",
    )

    add_review_appendix(
        doc,
        [
            "Have counsel and the insurer review the activity scope, risk description, emergency wording and liability allocation against actual services and cover.",
            "Confirm the secure pre-participation/onboarding process for relevant medical or accessibility information, including Article 9 basis, access, retention and staff response.",
            "Confirm staff qualifications, first-aid arrangements, emergency action plan, equipment inspection and incident reporting; the document cannot substitute for operational controls.",
            "Confirm that checkout and support operations continue to reject third-party purchase of an adult membership.",
            "Implement immutable or append-only acceptance evidence: exact version, statements, typed name, verified context and timestamp; do not rely on browser localStorage or a mutable profile field.",
            "Keep optional media consent outside this document and outside membership eligibility.",
            "Review accessibility, Welsh-language needs and a paper/assisted alternative for anyone unable to use the electronic flow.",
        ],
        [
            ("Unfair Contract Terms Act 1977, section 2", "https://www.legislation.gov.uk/ukpga/1977/50/section/2"),
            ("Consumer Rights Act 2015", "https://www.legislation.gov.uk/ukpga/2015/15/contents"),
            ("CMA unfair contract terms guidance", "https://www.gov.uk/government/publications/unfair-contract-terms-cma37"),
            ("Law Commission: electronic execution of documents", "https://lawcom.gov.uk/project/electronic-execution-of-documents/"),
            ("HSE health and safety basics for leisure activities", "https://www.hse.gov.uk/entertainment/leisure/basics.htm"),
            ("ICO special-category data guidance", "https://ico.org.uk/for-organisations/uk-gdpr-guidance-and-resources/lawful-basis/special-category-data/"),
        ],
        high_priority=(
            "A waiver does not make an unsafe system safe and cannot exclude liability for death or personal injury caused by negligence. Legal and insurance review must be paired with the real risk assessment, coaching controls, equipment maintenance, emergency plan and incident process."
        ),
    )
    path = OUTPUT_DIR / spec.filename
    doc.save(path)
    return path


def build_guardian() -> Path:
    spec = SPECS["guardian"]
    doc = Document()
    configure_document(doc, spec)
    add_document_title(doc, spec)
    add_draft_banner(doc)

    add_callout(
        doc,
        "Guardian is the payer and signer",
        "The adult guardian signs this addendum, accepts the Membership Terms and pays. The child receives an age-appropriate explanation during onboarding. This document does not ask a child to waive rights that cannot lawfully be waived.",
        kind="info",
    )

    doc.add_heading("1. Youth membership covered", level=1)
    add_para(
        doc,
        "This addendum applies separately to every named child enrolled in Youngstars (minimum age 6, maximum age 11) or Teenstars (minimum age 12, maximum age 16). One Checkout may cover between one and ten children in one selected plan only; Youngstars and Teenstars cannot be combined in one Checkout. References below to ‘the child’ apply separately to each listed child. A guardian must contact us if any date of birth or selected plan is wrong.",
    )
    add_para(
        doc,
        "The initial youth membership does not provide the child with AlphaWOD access. Completing Stripe Checkout forms the membership contract but does not reserve a first session. During the founding presale, nothing is charged today; membership starts and the first monthly payment is taken on 1 September 2026. Zero Alpha Fitness will contact the guardian by email to arrange onboarding and the first session.",
    )

    doc.add_heading("2. Guardian authority and information", level=1)
    add_para(
        doc,
        "I confirm that I am aged 18 or over, I am each named child’s parent or legal guardian or otherwise have lawful authority to make the decisions and commitments in this addendum for every named child, and I am the payer. If responsibility is shared or restricted by a court order or other arrangement, I confirm that signing and enrolling each child is permitted and I will tell Zero Alpha Fitness promptly about any relevant restriction or change.",
    )
    add_para(
        doc,
        "I confirm that each child’s name, date of birth and selected plan, and my relationship and contact details for each child, are accurate. I understand that Zero Alpha Fitness may pause onboarding or participation while it reasonably verifies eligibility or authority.",
    )

    doc.add_heading("3. Activities and risks", level=1)
    add_para(
        doc,
        "Youth activities may include age-appropriate functional fitness, movement skills, bodyweight and resistance exercises, conditioning, games, coached circuits, use of suitable equipment and related warm-up or cool-down activities. The programme should be adapted to age, maturity, experience and the session plan.",
    )
    add_para(
        doc,
        "Physical activity has inherent risks even where reasonable care is taken. These may include slips, trips, falls, collisions, overexertion, soreness, strains, sprains, fractures, equipment-related injury, aggravation of an existing condition and, rarely, serious or permanent injury. I understand those risks and consent to the child’s participation subject to Zero Alpha Fitness’s duty to use reasonable care, appropriate safeguarding and age-appropriate supervision.",
    )

    doc.add_heading("4. Child’s readiness and safety information", level=1)
    add_bullets(
        doc,
        [
            "I will consider whether the child is well enough and ready to participate, seek medical advice where appropriate, and follow professional advice.",
            "I will not send the child to participate while acutely ill, injured, impaired or subject to advice that makes participation unsafe.",
            "I will tell Zero Alpha Fitness promptly, through the secure onboarding route, about information reasonably necessary to adapt or safely manage participation.",
            "I will keep relevant information current and will not ask the child to conceal a material safety concern.",
        ],
    )
    add_callout(
        doc,
        "Do not enter health details in checkout",
        "The public purchase, signature and support fields are not the place for a child’s medical information. Use only the separate secure onboarding channel and read its just-in-time privacy information.",
        kind="warning",
    )

    doc.add_heading("5. Safeguarding, supervision and handover", level=1)
    add_bullets(
        doc,
        [
            "I and the child will follow reasonable safeguarding, behaviour, clothing, equipment and safety instructions.",
            "I will comply with the stated arrival, sign-in, handover and collection process and will provide information about any person authorised or prohibited from collecting the child where lawfully required.",
            "I will arrive and collect on time and will not assume supervision starts before handover or continues after the stated collection point.",
            "I understand that age-appropriate supervised activity still requires the child to listen, behave safely and tell a coach if they feel pain, unwell, unsafe or unable to continue.",
            "Zero Alpha Fitness may stop or adapt an activity, contact me, or take proportionate immediate action where reasonably necessary for safety or safeguarding.",
        ],
    )

    doc.add_heading("6. Emergency assistance", level=1)
    add_para(
        doc,
        "If staff reasonably believe urgent assistance is needed and cannot contact me in time, I authorise them to provide or arrange proportionate first aid, contact emergency services and share the minimum information reasonably necessary to protect the child. Staff will try to contact me as soon as reasonably practicable. This does not guarantee a particular treatment and does not replace professional medical advice.",
    )

    doc.add_heading("7. Membership, payment and cancellation", level=1)
    add_para(
        doc,
        "As payer, I accept the standard monthly price of £30 per Youngstars child or £35 per Teenstars child and the recurring authority. When the contracted quantity is two or more children, a recurring 15% multi-child discount applies to the full same-plan subtotal. It does not apply at quantity one. Stripe shows the quantity, price per child, undiscounted subtotal, discount and resulting recurring total before confirmation. Non-attendance does not change the contracted quantity, total or discount. During the founding presale, Stripe shows £0 due today and the first payment on 1 September 2026; after opening, Stripe may show immediate proration to the next first. I also accept the three-day past-due rule after service has started, the rule that a failed first scheduled payment means membership and access do not start, the no-pause rule, and the ordinary 14-day pre-renewal cancellation deadline. I understand that statutory cooling-off and refund rights remain separate and cannot be removed. The Membership Terms and Cancellation, Refund and Cooling-off Policy explain the details.",
    )
    add_para(
        doc,
        "The youth Checkout creates one subscription covering all listed children. An ordinary or cooling-off cancellation ends the whole subscription and all listed children’s places; the online flow does not cancel one child separately. A requested addition, removal, plan, price or discount change takes effect only when Zero Alpha Fitness confirms it in writing, including the effective date and resulting recurring total, and may require cancellation and a new Checkout.",
    )
    add_para(
        doc,
        "Reaching the top of an age band does not automatically move a child, remove a place or change the recurring total. Zero Alpha Fitness will contact the payer about any transition. No plan, contracted-quantity, price or discount change takes effect until confirmed in writing, and a new Checkout may be required.",
    )

    doc.add_heading("8. Conduct and proportionate restriction", level=1)
    add_para(
        doc,
        "Zero Alpha Fitness may adapt, pause or end the child’s participation in a session where reasonably necessary for safety, safeguarding, serious disruption or suitability. Any longer restriction or membership termination will be handled fairly under the Membership Terms, with an explanation and review opportunity where appropriate. Reasonable adjustments and relevant legal duties will be considered.",
    )

    doc.add_heading("9. Liability and the child’s rights", level=1)
    add_para(
        doc,
        "Zero Alpha Fitness must provide services with reasonable care and skill. Nothing in this addendum excludes or limits liability for death or personal injury caused by negligence, fraud or fraudulent misrepresentation, breach of statutory rights, safeguarding duties, or any other liability that cannot lawfully be excluded or limited.",
    )
    add_para(
        doc,
        "I acknowledge inherent risks that remain despite reasonable care. So far as fair and lawful, Zero Alpha Fitness is not responsible for harm caused by deliberate unsafe conduct, material breach of clear instructions, or materially inaccurate information knowingly supplied by the guardian, taking account of the child’s age and all circumstances. The child retains their own rights.",
    )

    doc.add_heading("10. Privacy and the child’s voice", level=1)
    add_para(
        doc,
        "The Privacy Notice explains how child, guardian, signature, membership and incident information is used. I will make the Notice available to the child and support Zero Alpha Fitness in giving an age-appropriate explanation. Data-protection rights belong to the child. My ability to exercise them for the child depends on authority, capacity and the child’s best interests.",
    )
    add_para(
        doc,
        "This addendum is not consent to marketing, photography, video or promotional use. Any media permission will be optional, specific and separate, and refusal will not affect membership.",
    )

    doc.add_heading("11. Electronic signature and changes", level=1)
    add_para(
        doc,
        "I intend my typed name and affirmative submission to authenticate and sign this addendum electronically. I will be able to review it before signing and receive or access a durable copy afterwards. The record will identify the document version and time of acceptance.",
    )
    continuing_effect = add_para(
        doc,
        "This addendum continues while the child participates. A material change requires clear notice and any fresh agreement reasonably or legally required. If a provision is unenforceable, the rest continues so far as lawful. The Membership Terms’ law and jurisdiction wording applies, subject to the child’s and consumer’s mandatory rights.",
    )
    continuing_effect.paragraph_format.keep_together = True

    doc.add_heading("Guardian electronic signing arrangement", level=1)
    add_quote(doc, "☐ I confirm that I am aged 18 or over, I am each named child’s parent or legal guardian or otherwise have lawful authority to enrol every named child, and I am the payer.")
    add_quote(doc, "☐ I have read and agree to the Parent/Guardian Consent and Youth Membership Addendum for each named child. I understand the activities and inherent risks and consent to each child’s participation, subject to their statutory rights and Zero Alpha Fitness’s duty to take reasonable care.")
    add_para(
        doc,
        "Required fields: every child’s full name, date of birth and selected-plan eligibility; contracted quantity; price per child; undiscounted subtotal; discount rate and amount; resulting recurring total; guardian full legal name; relationship and authority for each child; payer email collected by Stripe or the account-claim flow; typed signature; the exact acceptance statements; signature date/time; membership/order reference; addendum document ID and version.",
    )

    add_review_appendix(
        doc,
        [
            "Have safeguarding lead, counsel and insurer review the age bands, programme scope, coach-to-child supervision, handover/collection and emergency wording against actual operations.",
            "Define the age transition rule when a Youngstars child turns 12 or a Teenstars child turns 17, including notice, pricing, service transition and renewed documents.",
            "Confirm what evidence of guardian authority may be requested, how shared responsibility/court restrictions are handled, and who may collect a child.",
            "Complete a child-focused DPIA and legitimate-interests assessment; prepare age-appropriate notices for ages 6–11 and 12–16.",
            "Design the secure safety-information route and Article 9 process before requesting any child health information.",
            "Confirm youth incident reporting, first aid, emergency contact and escalation procedures. The purchase flow currently does not collect a phone number; assess whether a separate onboarding emergency contact is operationally and legally required.",
            "Implement immutable or append-only evidence of the guardian declaration, exact terms, typed signature, verified identity context and timestamp.",
            "Keep media permission separate and optional; confirm photography controls in the safeguarding policy.",
            "Review accessibility, assisted-signing and Welsh-language needs for guardians and children.",
        ],
        [
            ("Consumer Rights Act 2015", "https://www.legislation.gov.uk/ukpga/2015/15/contents"),
            ("Unfair Contract Terms Act 1977, section 2", "https://www.legislation.gov.uk/ukpga/1977/50/section/2"),
            ("Children Act 1989", "https://www.legislation.gov.uk/ukpga/1989/41/contents"),
            ("HSE statement on children’s play and leisure", "https://www.hse.gov.uk/entertainment/childs-play-statement.htm"),
            ("ICO children and the UK GDPR", "https://ico.org.uk/for-organisations/uk-gdpr-guidance-and-resources/childrens-information/children-and-the-uk-gdpr/"),
            ("ICO children’s data-protection rights", "https://ico.org.uk/for-organisations/uk-gdpr-guidance-and-resources/childrens-information/children-and-the-uk-gdpr/what-data-protection-rights-do-children-have/"),
            ("Law Commission: electronic execution of documents", "https://lawcom.gov.uk/project/electronic-execution-of-documents/"),
        ],
        high_priority=(
            "The youth checkout cannot safely launch on contract text alone. The real safeguarding, age-transition, handover/collection, emergency-contact and special-category-data processes must be defined, reviewed and matched to this addendum."
        ),
        compact=True,
    )
    path = OUTPUT_DIR / spec.filename
    doc.save(path)
    return path


def paragraph_plain_text(paragraph: Paragraph) -> str:
    """Return visible paragraph text, including text inside hyperlinks."""
    parts: list[str] = []
    for node in paragraph._p.iter():
        if node.tag == qn("w:t") and node.text:
            parts.append(node.text)
        elif node.tag in {qn("w:br"), qn("w:cr")}:
            parts.append("\n")
        elif node.tag == qn("w:tab"):
            parts.append("\t")
    return "".join(parts).strip()


def iter_document_blocks(doc: DocumentObject):
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def document_plain_text(doc: DocumentObject) -> str:
    """Create deterministic customer-facing UTF-8 text from the final DOCX."""
    blocks: list[str] = []
    number_index = 0
    for block in iter_document_blocks(doc):
        if isinstance(block, Paragraph):
            text = paragraph_plain_text(block)
            if not text:
                continue
            style_name = block.style.name if block.style is not None else ""
            if style_name == "List Bullet":
                text = f"• {text}"
                number_index = 0
            elif style_name == "List Number":
                number_index += 1
                text = f"{number_index}. {text}"
            else:
                number_index = 0
            blocks.append(text)
            continue

        number_index = 0
        rows: list[str] = []
        for row in block.rows:
            cells: list[str] = []
            for cell in row.cells:
                paragraphs = [
                    paragraph_plain_text(paragraph)
                    for paragraph in cell.paragraphs
                ]
                cell_text = "\n\n".join(text for text in paragraphs if text)
                if cell_text:
                    cells.append(cell_text)
            if cells:
                rows.append("\n\n".join(cells))
        if rows:
            blocks.append("\n\n".join(rows))

    text = "\n\n".join(blocks).strip() + "\n"
    forbidden = (
        "Draft for owner, legal and operational review",
        "not approved for publication",
        "Legal review appendix",
        "Internal review material",
        "Items to confirm",
        "Primary legal and regulatory sources",
    )
    for marker in forbidden:
        if marker.lower() in text.lower():
            raise ValueError(f"Final customer copy still contains review marker: {marker}")
    return text


def write_plain_text_bundle(paths: Sequence[Path], output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    specs_by_filename = {spec.filename: spec for spec in SPECS.values()}
    manifest: dict[str, object] = {
        "effectiveDate": "2026-08-23",
        "hashCovers": "UTF-8 bytes of content",
        "documents": {},
    }
    documents = manifest["documents"]
    assert isinstance(documents, dict)

    for path in paths:
        spec = specs_by_filename[path.name]
        content = document_plain_text(Document(path))
        target = output_dir / f"{spec.doc_id}.txt"
        target.write_text(content, encoding="utf-8", newline="\n")
        encoded = content.encode("utf-8")
        documents[spec.registry_key] = {
            "title": spec.title,
            "version": spec.doc_id,
            "effectiveDate": "2026-08-23",
            "filename": target.name,
            "sha256": hashlib.sha256(encoded).hexdigest(),
            "bytes": len(encoded),
        }

    manifest_path = output_dir / "manifest.json"
    manifest_path.write_text(
        json.dumps(manifest, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    return manifest_path


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Build final customer-facing Zero Alpha Fitness legal documents."
    )
    parser.add_argument(
        "--plain-text-output-dir",
        type=Path,
        help="Also emit canonical customer-facing UTF-8 text and a hash manifest.",
    )
    args = parser.parse_args()

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    paths = [
        build_terms(),
        build_privacy(),
        build_cancellation(),
        build_adult_waiver(),
        build_guardian(),
    ]
    for path in paths:
        print(path.resolve())
    if args.plain_text_output_dir is not None:
        manifest_path = write_plain_text_bundle(paths, args.plain_text_output_dir)
        print(manifest_path.resolve())


if __name__ == "__main__":
    main()
